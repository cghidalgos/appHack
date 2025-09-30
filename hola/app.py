
import os
from flask import Flask, render_template, request, redirect, url_for, session, flash
from pymongo import MongoClient

app = Flask(__name__)
app.secret_key = 'supersecretkey'
UPLOAD_FOLDER = 'uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# Cargar variables de entorno desde .env y mostrar la API key
from dotenv import load_dotenv
load_dotenv()
print("API KEY:", os.getenv('OPENAI_API_KEY'))

# Conexión a MongoDB
MONGO_URI = "mongodb+srv://cgiohidalgo:holamundo123@giovanny.qlddw6j.mongodb.net/?retryWrites=true&w=majority&appName=giovanny"
mongo_client = MongoClient(MONGO_URI)
db = mongo_client['giovanny']
usuarios_col = db['usuarios']
historias_col = db['historias']

@app.route('/')
def landing():
    return render_template('landing.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        if password == '12345':
            session['username'] = username
            session['role'] = request.form['role']
            if session['role'] == 'admin':
                return redirect(url_for('admin_dashboard'))
            else:
                return redirect(url_for('user_dashboard'))
        else:
            flash('Contraseña incorrecta')
    return render_template('login.html')

@app.route('/admin')
def admin_dashboard():
    if session.get('role') != 'admin':
        return redirect(url_for('login'))
    return render_template('admin_dashboard.html')

@app.route('/user')
def user_dashboard():
    if session.get('role') != 'user':
        return redirect(url_for('login'))
    return render_template('user_dashboard.html')

# Ruta para listar todos los usuarios y ver su información
@app.route('/admin/all_users')
def admin_all_users():
    if session.get('role') != 'admin':
        return redirect(url_for('login'))
    usuarios = list(usuarios_col.find())
    for usuario in usuarios:
        usuario['historias'] = list(historias_col.find({'usuario_id': usuario['_id']}))
    return render_template('admin_all_users.html', usuarios=usuarios)

# Ruta para cargar documento y extraer información
from werkzeug.utils import secure_filename
import openai
from PyPDF2 import PdfReader
from docx import Document
from PIL import Image

openai.api_key = os.getenv('OPENAI_API_KEY')

def extract_text_from_file(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.pdf':
        reader = PdfReader(filepath)
        text = "\n".join(page.extract_text() or '' for page in reader.pages)
        return text
    elif ext in ['.doc', '.docx']:
        doc = Document(filepath)
        return "\n".join([p.text for p in doc.paragraphs])
    elif ext in ['.png', '.jpg', '.jpeg']:
        try:
            import pytesseract
            img = Image.open(filepath)
            return pytesseract.image_to_string(img)
        except Exception:
            return "No se pudo extraer texto de la imagen."
    return ""

def extract_info_with_openai(text):
    prompt = (
        "Extrae la siguiente información de la historia clínica en formato JSON, usando exactamente estos campos: "
        "nombre, cedula, edad, diagnostico, antecedentes, tratamientos, fecha. "
        "Si algún dato no está presente, pon el valor como null.\n\nHistoria clínica:\n\n" + text + "\n\nEjemplo de respuesta:\n{\n  'nombre': 'Juan Perez',\n  'cedula': '123456789',\n  'edad': 45,\n  'diagnostico': 'Hipertensión',\n  'antecedentes': 'Ninguno',\n  'tratamientos': 'Enalapril',\n  'fecha': '2023-09-30'\n}"
    )
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}]
        )
        return response.choices[0].message['content']
    except Exception as e:
        return f"Error al extraer información: {e}"

@app.route('/admin/upload', methods=['GET', 'POST'])
def admin_upload():
    if session.get('role') != 'admin':
        return redirect(url_for('login'))
    if request.method == 'POST':
        file = request.files['file']
        if file:
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            text = extract_text_from_file(filepath)
            print("Texto extraído:", text[:500])
            info = extract_info_with_openai(text)
            print("Respuesta OpenAI:", info)
            import json, re
            nombre, cedula = None, None
            datos_extraidos = {}
            try:
                json_match = re.search(r'\{.*\}', info, re.DOTALL)
                if json_match:
                    info_json = json_match.group(0)
                    info_json = info_json.replace("'", '"')
                    datos_extraidos = json.loads(info_json)
                    nombre = datos_extraidos.get('nombre')
                    cedula = datos_extraidos.get('cedula')
            except Exception as e:
                print("Error al parsear JSON:", e)
            if not (nombre and cedula):
                match = re.search(r"nombre[:\s]*([\w\s]+).*cedula[:\s]*([\w\d]+)", info, re.I|re.S)
                if match:
                    nombre = match.group(1).strip()
                    cedula = match.group(2).strip()
            if not (nombre and cedula):
                for line in info.splitlines():
                    if 'nombre' in line.lower():
                        nombre = line.split(':')[-1].strip()
                    if 'cedula' in line.lower():
                        cedula = line.split(':')[-1].strip()
            if nombre and cedula:
                user = usuarios_col.find_one({'cedula': cedula})
                if not user:
                    user_id = usuarios_col.insert_one({'nombre': nombre, 'cedula': cedula}).inserted_id
                else:
                    user_id = user['_id']
                historias_col.insert_one({
                    'usuario_id': user_id,
                    'contenido': datos_extraidos if datos_extraidos else {'raw': info},
                    'archivo': filename
                })
                flash(f'Documento cargado y datos extraídos correctamente. Usuario: {nombre}, Cédula: {cedula}')
                return redirect(url_for('admin_upload'))
            else:
                flash('No se pudo extraer nombre y cédula. Revisa el formato del documento o la respuesta de OpenAI. Respuesta: {}'.format(info))
                return redirect(url_for('admin_upload'))
    return render_template('admin_upload.html')

# Ruta para buscar usuario (admin)
@app.route('/admin/search', methods=['GET'])
def admin_search():
    if session.get('role') != 'admin':
        return redirect(url_for('login'))
    cedula = request.args.get('cedula')
    nombre = request.args.get('nombre')
    if not cedula and not nombre:
        usuarios = list(usuarios_col.find())
        for usuario in usuarios:
            usuario['historias'] = list(historias_col.find({'usuario_id': usuario['_id']}))
        return render_template('admin_all_users.html', usuarios=usuarios)
    usuario = None
    if cedula:
        usuario = usuarios_col.find_one({'cedula': cedula})
    elif nombre:
        usuario = usuarios_col.find_one({'nombre': {'$regex': nombre, '$options': 'i'}})
    historias = []
    if usuario:
        historias = list(historias_col.find({'usuario_id': usuario['_id']}))
    return render_template('admin_result.html', usuario=usuario, historias=historias)

# Ruta para buscar historia clínica (usuario)
@app.route('/user/search', methods=['GET'])
def user_search():
    if session.get('role') != 'user':
        return redirect(url_for('login'))
    cedula = request.args.get('cedula')
    nombre = request.args.get('nombre')
    usuario = None
    if cedula:
        usuario = usuarios_col.find_one({'cedula': cedula})
    elif nombre:
        usuario = usuarios_col.find_one({'nombre': {'$regex': nombre, '$options': 'i'}})
    historias = []
    if usuario:
        historias = list(historias_col.find({'usuario_id': usuario['_id']}))
    return render_template('user_result.html', usuario=usuario, historias=historias)

if __name__ == '__main__':
    app.run(debug=True, host="0.0.0.0")
import os
from flask import Flask, render_template, request, redirect, url_for, session, flash
from pymongo import MongoClient

app = Flask(__name__)
app.secret_key = 'supersecretkey'
UPLOAD_FOLDER = 'uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# Cargar variables de entorno desde .env y mostrar la API key
from dotenv import load_dotenv
load_dotenv()
print("API KEY:", os.getenv('OPENAI_API_KEY'))

# Conexión a MongoDB
MONGO_URI = "mongodb+srv://cgiohidalgo:holamundo123@giovanny.qlddw6j.mongodb.net/?retryWrites=true&w=majority&appName=giovanny"
mongo_client = MongoClient(MONGO_URI)
db = mongo_client['giovanny']
usuarios_col = db['usuarios']
historias_col = db['historias']

@app.route('/')
def landing():
    return render_template('landing.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        if password == '12345':
            session['username'] = username
            session['role'] = request.form['role']
            if session['role'] == 'admin':
                return redirect(url_for('admin_dashboard'))
            else:
                return redirect(url_for('user_dashboard'))
        else:
            flash('Contraseña incorrecta')
    return render_template('login.html')

@app.route('/admin')
def admin_dashboard():
    if session.get('role') != 'admin':
        return redirect(url_for('login'))
    return render_template('admin_dashboard.html')

@app.route('/user')
def user_dashboard():
    if session.get('role') != 'user':
        return redirect(url_for('login'))
    return render_template('user_dashboard.html')

# Aquí irán las rutas para cargar documentos, extraer info, buscar usuarios, etc.
# Ruta para cargar documento y extraer información
from werkzeug.utils import secure_filename
import openai
from PyPDF2 import PdfReader
from docx import Document
from PIL import Image
from dotenv import load_dotenv

# Cargar variables de entorno desde .env
load_dotenv()
print("API KEY:", os.getenv('OPENAI_API_KEY'))
openai.api_key = os.getenv('OPENAI_API_KEY')

def extract_text_from_file(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.pdf':
        reader = PdfReader(filepath)
        text = "\n".join(page.extract_text() or '' for page in reader.pages)
        return text
    elif ext in ['.doc', '.docx']:
        doc = Document(filepath)
        return "\n".join([p.text for p in doc.paragraphs])
    elif ext in ['.png', '.jpg', '.jpeg']:
        try:
            import pytesseract
            img = Image.open(filepath)
            return pytesseract.image_to_string(img)
        except Exception:
            return "No se pudo extraer texto de la imagen."
    return ""

def extract_info_with_openai(text):
    prompt = (
        "Extrae la siguiente información de la historia clínica en formato JSON, usando exactamente estos campos: "
        "nombre, cedula, edad, diagnostico, antecedentes, tratamientos, fecha. "
        "Si algún dato no está presente, pon el valor como null.\n\nHistoria clínica:\n\n" + text + "\n\nEjemplo de respuesta:\n{\n  'nombre': 'Juan Perez',\n  'cedula': '123456789',\n  'edad': 45,\n  'diagnostico': 'Hipertensión',\n  'antecedentes': 'Ninguno',\n  'tratamientos': 'Enalapril',\n  'fecha': '2023-09-30'\n}"
    )
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}]
        )
        return response.choices[0].message['content']
    except Exception as e:
        return f"Error al extraer información: {e}"

@app.route('/admin/upload', methods=['GET', 'POST'])
def admin_upload():
    if session.get('role') != 'admin':
        return redirect(url_for('login'))
    if request.method == 'POST':
        file = request.files['file']
        if file:
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            text = extract_text_from_file(filepath)
            print("Texto extraído:", text[:500])
            info = extract_info_with_openai(text)
            print("Respuesta OpenAI:", info)
            import json, re
            nombre, cedula = None, None
            datos_extraidos = {}
            # Intentar parsear JSON robusto
            try:
                # Buscar el primer bloque JSON en la respuesta
                json_match = re.search(r'\{.*\}', info, re.DOTALL)
                if json_match:
                    info_json = json_match.group(0)
                    # Reemplazar comillas simples por dobles y limpiar
                    info_json = info_json.replace("'", '"')
                    datos_extraidos = json.loads(info_json)
                    nombre = datos_extraidos.get('nombre')
                    cedula = datos_extraidos.get('cedula')
            except Exception as e:
                print("Error al parsear JSON:", e)
            # Fallback: buscar por regex si no se encontró en JSON
            if not (nombre and cedula):
                match = re.search(r"nombre[:\s]*([\w\s]+).*cedula[:\s]*([\w\d]+)", info, re.I|re.S)
                if match:
                    nombre = match.group(1).strip()
                    cedula = match.group(2).strip()
            # Fallback: buscar por líneas si sigue sin encontrarse
            if not (nombre and cedula):
                for line in info.splitlines():
                    if 'nombre' in line.lower():
                        nombre = line.split(':')[-1].strip()
                    if 'cedula' in line.lower():
                        cedula = line.split(':')[-1].strip()
            # Guardar y asociar si se encontró nombre y cédula
            if nombre and cedula:
                # Buscar usuario por cédula
                user = usuarios_col.find_one({'cedula': cedula})
                if not user:
                    # Crear usuario si no existe
                    user_id = usuarios_col.insert_one({'nombre': nombre, 'cedula': cedula}).inserted_id
                else:
                    user_id = user['_id']
                # Guardar historia clínica asociada
                historias_col.insert_one({
                    'usuario_id': user_id,
                    'contenido': datos_extraidos if datos_extraidos else {'raw': info},
                    'archivo': filename
                })
                flash(f'Documento cargado y datos extraídos correctamente. Usuario: {nombre}, Cédula: {cedula}')
                return redirect(url_for('admin_upload'))
            else:
                flash('No se pudo extraer nombre y cédula. Revisa el formato del documento o la respuesta de OpenAI. Respuesta: {}'.format(info))
                return redirect(url_for('admin_upload'))
    return render_template('admin_upload.html')

# Ruta para buscar usuario (admin)
@app.route('/admin/search', methods=['GET'])
def admin_search():
    if session.get('role') != 'admin':
        return redirect(url_for('login'))
    cedula = request.args.get('cedula')
    nombre = request.args.get('nombre')
    if not cedula and not nombre:
        # Mostrar todos los usuarios y sus historias
        usuarios = list(usuarios_col.find())
        for usuario in usuarios:
            usuario['historias'] = list(historias_col.find({'usuario_id': usuario['_id']}))
        return render_template('admin_all_users.html', usuarios=usuarios)
    usuario = None
    if cedula:
        usuario = usuarios_col.find_one({'cedula': cedula})
    elif nombre:
        usuario = usuarios_col.find_one({'nombre': {'$regex': nombre, '$options': 'i'}})
    historias = []
    if usuario:
        historias = list(historias_col.find({'usuario_id': usuario['_id']}))
    return render_template('admin_result.html', usuario=usuario, historias=historias)

# Ruta para buscar historia clínica (usuario)
@app.route('/user/search', methods=['GET'])
def user_search():
    if session.get('role') != 'user':
        return redirect(url_for('login'))
    cedula = request.args.get('cedula')
    nombre = request.args.get('nombre')
    usuario = None
    if cedula:
        usuario = usuarios_col.find_one({'cedula': cedula})
    elif nombre:
        usuario = usuarios_col.find_one({'nombre': {'$regex': nombre, '$options': 'i'}})
    historias = []
    if usuario:
        historias = list(historias_col.find({'usuario_id': usuario['_id']}))
    return render_template('user_result.html', usuario=usuario, historias=historias)

if __name__ == '__main__':
    app.run(debug=True, host="0.0.0.0")
